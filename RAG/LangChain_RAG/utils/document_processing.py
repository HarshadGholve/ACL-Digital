import os
from bs4 import BeautifulSoup
from langchain_community.document_loaders import WebBaseLoader
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from PyPDF2 import PdfReader
from docx import Document as DocxDocument


def process_uploaded_file(file_path: str) -> list[Document]:
    ext = os.path.splitext(file_path)[1].lower()
    documents = []

    if ext == ".pdf":
        reader = PdfReader(file_path)
        documents = [Document(page_content=page.extract_text()) for page in reader.pages]
    elif ext == ".docx":
        doc = DocxDocument(file_path)
        documents = [Document(page_content=para.text) for para in doc.paragraphs]
    elif ext == ".txt":
        with open(file_path, "r") as f:
            documents = [Document(page_content=f.read())]
    return documents


def process_url(url: str) -> list[Document]:
    loader = WebBaseLoader(
        web_paths=(url,),
        bs_kwargs=dict(parse_only=BeautifulSoup.SoupStrainer())
    )
    return loader.load()


def split_documents(documents: list[Document], chunk_size: int = 1000, chunk_overlap: int = 200) -> list[Document]:
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    return text_splitter.split_documents(documents)


def add_documents_to_store(vector_store, documents: list[Document]):
    vector_store.add_documents(documents)
